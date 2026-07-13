from __future__ import annotations

from pathlib import Path

from docx import Document

from .base import BaseParser, ParsedFile


class DocxParser(BaseParser):
    supported_suffixes = (".docx",)
    kind = "document"

    def parse(self, path: Path) -> ParsedFile:
        doc = Document(str(path))
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        tables = []
        for table_index, table in enumerate(doc.tables, 1):
            rows = ["\t".join(cell.text.strip() for cell in row.cells) for row in table.rows]
            table_text = "\n".join(row for row in rows if row.strip())
            if table_text:
                tables.append(f"[table {table_index}]\n{table_text}")
        text = "\n".join([*paragraphs, *tables])
        return ParsedFile(
            path=path,
            kind=self.kind,
            text_summary=text if text else "空文档，或仅包含图片/形状等非文本元素。",
            metadata={"paragraph_count": len(paragraphs), "table_count": len(doc.tables), "chars": len(text)},
        )
