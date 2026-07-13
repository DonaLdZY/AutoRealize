from __future__ import annotations

import hashlib
import re
from dataclasses import dataclass
from pathlib import Path
from typing import Any, Iterable

from docx import Document
from pypdf import PdfReader

from .utils.filesystem import rel, walk_files
from .utils.safe_json import write_json_safe


DOCUMENT_SUFFIXES = {".pdf", ".docx", ".txt", ".md", ".rst", ".log"}


@dataclass(frozen=True)
class DocumentChunk:
    chunk_id: str
    document_id: str
    source_file: str
    locator: str
    ordinal: int
    text: str

    def as_dict(self, *, include_text: bool = True) -> dict[str, Any]:
        payload: dict[str, Any] = {
            "chunk_id": self.chunk_id,
            "document_id": self.document_id,
            "source_file": self.source_file,
            "locator": self.locator,
            "ordinal": self.ordinal,
            "chars": len(self.text),
        }
        if include_text:
            payload["text"] = self.text
        return payload


class LocalDocumentIndex:
    """Deterministic full-text store used by the QDI live context."""

    def __init__(self, root: Path, documents: list[dict[str, Any]], chunks: list[DocumentChunk]) -> None:
        self.root = root
        self.documents = documents
        self.chunks = chunks
        self._by_chunk = {chunk.chunk_id: chunk for chunk in chunks}
        self._ordinals: dict[str, list[DocumentChunk]] = {}
        for chunk in chunks:
            self._ordinals.setdefault(chunk.document_id, []).append(chunk)
        for values in self._ordinals.values():
            values.sort(key=lambda item: item.ordinal)

    @classmethod
    def build(
        cls,
        *,
        data_root: Path,
        store_root: Path,
        chunk_chars: int = 2200,
        chunk_overlap_chars: int = 200,
    ) -> "LocalDocumentIndex":
        store_root.mkdir(parents=True, exist_ok=True)
        documents: list[dict[str, Any]] = []
        all_chunks: list[DocumentChunk] = []
        for path in walk_files(data_root):
            if path.suffix.lower() not in DOCUMENT_SUFFIXES:
                continue
            source = rel(path, data_root)
            document_id = "doc_" + hashlib.sha256(source.encode("utf-8", errors="replace")).hexdigest()[:16]
            try:
                sections, metadata = _extract_document_sections(path)
            except Exception as exc:  # One malformed document must not disable QDI retrieval.
                documents.append(
                    {
                        "document_id": document_id,
                        "source_file": source,
                        "suffix": path.suffix.lower(),
                        "chunk_count": 0,
                        "text_chars": 0,
                        "parse_error": str(exc)[:500],
                    }
                )
                continue
            chunks: list[DocumentChunk] = []
            ordinal = 0
            for locator, text in sections:
                for part in _split_text(text, chunk_chars=chunk_chars, overlap_chars=chunk_overlap_chars):
                    ordinal += 1
                    chunks.append(
                        DocumentChunk(
                            chunk_id=f"{document_id}_c{ordinal:05d}",
                            document_id=document_id,
                            source_file=source,
                            locator=locator,
                            ordinal=ordinal,
                            text=part,
                        )
                    )
            document = {
                "document_id": document_id,
                "source_file": source,
                "suffix": path.suffix.lower(),
                "chunk_count": len(chunks),
                "text_chars": sum(len(chunk.text) for chunk in chunks),
                **metadata,
            }
            if path.suffix.lower() == ".pdf" and not chunks:
                document["requires_ocr"] = True
            documents.append(document)
            all_chunks.extend(chunks)
            write_json_safe(
                store_root / f"{document_id}.json",
                {"document": document, "chunks": [chunk.as_dict() for chunk in chunks]},
                indent=2,
            )
        write_json_safe(store_root / "manifest.json", {"documents": documents}, indent=2)
        return cls(store_root, documents, all_chunks)

    def manifest_for_prompt(self) -> list[dict[str, Any]]:
        return [
            {
                "document_id": item.get("document_id"),
                "source_file": item.get("source_file"),
                "suffix": item.get("suffix"),
                "pages": item.get("pages"),
                "paragraphs": item.get("paragraphs"),
                "chunk_count": item.get("chunk_count"),
                "text_chars": item.get("text_chars"),
                "requires_ocr": item.get("requires_ocr"),
                "parse_error": item.get("parse_error"),
                "full_text_local": True,
            }
            for item in self.documents
        ]

    def search(
        self,
        query: str,
        *,
        document_ids: Iterable[str] = (),
        source_files: Iterable[str] = (),
        top_k: int = 5,
        excerpt_chars: int = 1800,
    ) -> dict[str, Any]:
        query = str(query or "").strip()
        allowed_ids = {str(value).strip() for value in document_ids if str(value).strip()}
        allowed_sources = {_norm_path(value) for value in source_files if str(value).strip()}
        candidates = [
            chunk
            for chunk in self.chunks
            if (not allowed_ids or chunk.document_id in allowed_ids)
            and (not allowed_sources or _norm_path(chunk.source_file) in allowed_sources)
        ]
        ranked = sorted(
            ((self._score(query, chunk), chunk) for chunk in candidates),
            key=lambda item: (item[0], -item[1].ordinal),
            reverse=True,
        )
        matches = []
        for score, chunk in ranked:
            if score <= 0:
                continue
            payload = chunk.as_dict(include_text=False)
            payload.update(score=round(score, 6), excerpt=_query_centered_excerpt(chunk.text, query, excerpt_chars))
            matches.append(payload)
            if len(matches) >= max(1, int(top_k)):
                break
        return {
            "mode": "search",
            "query": query,
            "searched_chunks": len(candidates),
            "matches": matches,
            "no_match": not matches,
        }

    def read_chunks(
        self,
        chunk_ids: Iterable[str],
        *,
        neighbor_count: int = 0,
        max_chunks: int = 8,
        max_chars: int = 12000,
    ) -> dict[str, Any]:
        selected: dict[str, DocumentChunk] = {}
        for chunk_id in chunk_ids:
            chunk = self._by_chunk.get(str(chunk_id).strip())
            if chunk is None:
                continue
            selected[chunk.chunk_id] = chunk
            siblings = self._ordinals.get(chunk.document_id, [])
            start = max(0, chunk.ordinal - 1 - max(0, int(neighbor_count)))
            end = min(len(siblings), chunk.ordinal + max(0, int(neighbor_count)))
            for sibling in siblings[start:end]:
                selected[sibling.chunk_id] = sibling
        ordered = sorted(selected.values(), key=lambda item: (item.source_file, item.ordinal))[: max(1, int(max_chunks))]
        output = []
        visible_chars = 0
        truncated = False
        for chunk in ordered:
            remaining = max(0, int(max_chars) - visible_chars)
            if remaining <= 0:
                truncated = True
                break
            text = chunk.text[:remaining]
            payload = chunk.as_dict(include_text=False)
            payload["text"] = text
            payload["text_truncated"] = len(text) < len(chunk.text)
            output.append(payload)
            visible_chars += len(text)
            truncated = truncated or payload["text_truncated"]
        return {
            "mode": "read_chunks",
            "chunks": output,
            "requested_chunk_ids": [str(value) for value in chunk_ids],
            "visible_chars": visible_chars,
            "truncated": truncated or len(output) < len(ordered),
        }

    @staticmethod
    def _score(query: str, chunk: DocumentChunk) -> float:
        q = _normalize_text(query)
        text = _normalize_text(chunk.text)
        if not q or not text:
            return 0.0
        score = 8.0 * text.count(q)
        query_terms = _search_terms(q)
        if not query_terms:
            return score
        text_terms = set(_search_terms(text))
        overlap = sum(1 for term in query_terms if term in text_terms)
        score += 4.0 * overlap / max(1, len(set(query_terms)))
        score += sum(min(3, text.count(term)) * min(1.5, 0.25 + len(term) / 8) for term in set(query_terms))
        return score


def _extract_document_sections(path: Path) -> tuple[list[tuple[str, str]], dict[str, Any]]:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        reader = PdfReader(str(path))
        sections = []
        for index, page in enumerate(reader.pages, 1):
            text = (page.extract_text() or "").strip()
            if text:
                sections.append((f"page {index}", text))
        return sections, {"pages": len(reader.pages), "pages_extracted": len(sections)}
    if suffix == ".docx":
        document = Document(str(path))
        sections: list[tuple[str, str]] = []
        for index, paragraph in enumerate(document.paragraphs, 1):
            text = paragraph.text.strip()
            if text:
                sections.append((f"paragraph {index}", text))
        for table_index, table in enumerate(document.tables, 1):
            rows = ["\t".join(cell.text.strip() for cell in row.cells) for row in table.rows]
            text = "\n".join(row for row in rows if row.strip())
            if text:
                sections.append((f"table {table_index}", text))
        return sections, {"paragraphs": len(document.paragraphs), "tables": len(document.tables)}
    text = ""
    encoding = ""
    for candidate in ("utf-8-sig", "utf-8", "gb18030", "gbk", "latin-1"):
        try:
            text = path.read_text(encoding=candidate)
            encoding = candidate
            break
        except UnicodeDecodeError:
            continue
    return [("full text", text)] if text else [], {"encoding": encoding}


def _split_text(text: str, *, chunk_chars: int, overlap_chars: int) -> list[str]:
    clean = str(text or "").strip()
    if not clean:
        return []
    limit = max(400, int(chunk_chars))
    overlap = min(max(0, int(overlap_chars)), limit // 3)
    if len(clean) <= limit:
        return [clean]
    chunks = []
    start = 0
    while start < len(clean):
        end = min(len(clean), start + limit)
        if end < len(clean):
            split_at = max(clean.rfind("\n", start + limit // 2, end), clean.rfind("。", start + limit // 2, end))
            if split_at > start:
                end = split_at + 1
        chunks.append(clean[start:end].strip())
        if end >= len(clean):
            break
        start = max(start + 1, end - overlap)
    return [chunk for chunk in chunks if chunk]


def _normalize_text(value: Any) -> str:
    return re.sub(r"\s+", "", str(value or "").lower())


def _search_terms(value: str) -> list[str]:
    raw = str(value or "").lower()
    words = re.findall(r"[a-z0-9_]{2,}|[\u4e00-\u9fff]{2,}", raw)
    terms: list[str] = []
    for word in words:
        terms.append(word)
        if re.fullmatch(r"[\u4e00-\u9fff]+", word) and len(word) > 2:
            terms.extend(word[index : index + 2] for index in range(len(word) - 1))
    return terms[:200]


def _query_centered_excerpt(text: str, query: str, limit: int) -> str:
    if len(text) <= limit:
        return text
    normalized_query = str(query or "").strip()
    position = text.lower().find(normalized_query.lower()) if normalized_query else -1
    if position < 0:
        position = 0
    start = max(0, position - limit // 3)
    end = min(len(text), start + limit)
    prefix = "...[前文省略]..." if start else ""
    suffix = "...[后文省略]..." if end < len(text) else ""
    return prefix + text[start:end] + suffix


def _norm_path(value: Any) -> str:
    return str(value or "").replace("\\", "/").strip().lower()
